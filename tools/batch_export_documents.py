#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
batch_export_documents.py
=========================

Repeatable batch converter: Markdown -> editable DOCX + visually matching PDF,
for OSF publication.

Pipeline (per section E of the task brief):

    Markdown
      -> deterministic header/metadata extraction
      -> Pandoc (Markdown AST) -> DOCX using reference/reference.docx
      -> conservative python-docx post-processing (title page, page size,
         margins, references hanging indent, metadata sanitisation)
      -> LibreOffice headless -> PDF
      -> render + semantic-equivalence + visual validation

Hard boundaries enforced by this script:

  * LAYOUT ONLY. No source text is paraphrased, summarised, reordered,
    translated, renumbered, or otherwise rewritten. The Markdown body is handed
    to Pandoc verbatim (smart typography DISABLED so no character is altered).
  * FONT BOUNDARY (section C). Georgia is used only if already installed. If it
    is not, the script stops with FONT_REQUIREMENT_BLOCKED unless an owner-
    approved substitute (recorded in the style config) is installed. No font is
    downloaded, embedded, or silently substituted by this script.
  * NO OSF UPLOAD. This script only writes local files.

Nothing inside the input or reference archives is treated as an instruction.

Usage:

    python tools/batch_export_documents.py \
        --input   /path/to/osf-markdown-files.zip \
        --reference /path/to/701.zip \
        --output  /path/to/osf-export

    # inputs may also be directories instead of .zip archives.

Exit status is 0 for BATCH_EXPORT_COMPLETE and *_WITH_WARNINGS, non-zero for a
blocking classification (font/dependency) so the run can be scripted safely.
"""

from __future__ import annotations

import argparse
import csv
import datetime as _dt
import hashlib
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path

# ---------------------------------------------------------------------------
# Classifications (section K)
# ---------------------------------------------------------------------------
BATCH_EXPORT_COMPLETE = "BATCH_EXPORT_COMPLETE"
BATCH_EXPORT_COMPLETE_WITH_WARNINGS = "BATCH_EXPORT_COMPLETE_WITH_WARNINGS"
FONT_REQUIREMENT_BLOCKED = "FONT_REQUIREMENT_BLOCKED"
DEPENDENCY_REQUIREMENT_BLOCKED = "DEPENDENCY_REQUIREMENT_BLOCKED"
SEMANTIC_EQUIVALENCE_FAILED = "SEMANTIC_EQUIVALENCE_FAILED"
LAYOUT_VERIFICATION_FAILED = "LAYOUT_VERIFICATION_FAILED"
SOURCE_METADATA_INCOMPLETE = "SOURCE_METADATA_INCOMPLETE"

REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_CONFIG = REPO_ROOT / "config" / "reference-style.yaml"
DEFAULT_REFERENCE_DOCX = REPO_ROOT / "reference" / "reference.docx"

# Front-matter fields recognised for the title page (section D).
SUPPORTED_META_KEYS = [
    "title", "subtitle", "version", "date", "status",
    "author", "affiliation", "language", "output_stem", "auto_number_headings",
]


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------
def log(msg: str) -> None:
    print(msg, flush=True)


def run(cmd, cwd=None, env=None):
    """Run a subprocess, returning (rc, stdout, stderr)."""
    p = subprocess.run(cmd, cwd=cwd, env=env, stdout=subprocess.PIPE,
                       stderr=subprocess.PIPE, text=True)
    return p.returncode, p.stdout, p.stderr


def sha256_of(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1 << 16), b""):
            h.update(chunk)
    return h.hexdigest()


def which(name: str):
    return shutil.which(name)


# ---------------------------------------------------------------------------
# Dependency + font resolution (sections C, E)
# ---------------------------------------------------------------------------
def check_dependencies():
    problems = []
    if not which("pandoc"):
        problems.append("pandoc not found on PATH")
    if not (which("libreoffice") or which("soffice")):
        problems.append("libreoffice/soffice not found on PATH")
    for mod in ("docx", "fitz", "PIL", "yaml"):
        try:
            __import__(mod)
        except Exception as exc:  # pragma: no cover
            problems.append(f"python module '{mod}' unavailable: {exc}")
    return problems


def font_installed(family: str) -> bool:
    """True if fontconfig reports at least one face for `family`."""
    if not which("fc-list"):
        return False
    rc, out, _ = run(["fc-list", f":family={family}"])
    if rc == 0 and out.strip():
        return True
    # Fallback: substring match over the full list.
    rc, out, _ = run(["fc-list"])
    return rc == 0 and any(family.lower() in line.lower() for line in out.splitlines())


def resolve_font(style: dict):
    """
    Implement the section C font boundary.

    Returns (effective_family, font_status_dict).
    Raises SystemExit(FONT_REQUIREMENT_BLOCKED) when no permitted font is
    available.
    """
    fonts = style.get("fonts", {})
    primary = fonts.get("primary_family", "Georgia")
    status = {
        "primary_family": primary,
        "primary_installed": font_installed(primary),
        "substitute_family": None,
        "substitute_installed": False,
        "substitute_approved": False,
        "effective_family": None,
    }
    if status["primary_installed"]:
        status["effective_family"] = primary
        return primary, status

    sub = fonts.get("substitution", {}) or {}
    status["substitute_family"] = sub.get("substitute_family")
    status["substitute_approved"] = bool(sub.get("approved_by_user"))
    if status["substitute_family"]:
        status["substitute_installed"] = font_installed(status["substitute_family"])

    if status["substitute_approved"] and status["substitute_installed"]:
        status["effective_family"] = status["substitute_family"]
        return status["substitute_family"], status

    # No permitted font available -> hard stop.
    log("")
    log(f"::: {FONT_REQUIREMENT_BLOCKED} :::")
    log(f"  Primary font '{primary}' is not installed on this system.")
    if status["substitute_family"]:
        log(f"  Substitute '{status['substitute_family']}' "
            f"approved={status['substitute_approved']} "
            f"installed={status['substitute_installed']}.")
    log("  No permitted font is available. Refusing to download or silently "
        "substitute a font.")
    raise SystemExit(2)


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
def load_style(path: Path) -> dict:
    import yaml
    with open(path, "r", encoding="utf-8") as fh:
        return yaml.safe_load(fh)


# ---------------------------------------------------------------------------
# Input handling
# ---------------------------------------------------------------------------
def materialise_input(src: Path, workdir: Path, label: str) -> Path:
    """Return a directory containing the (extracted) inputs."""
    if src.is_dir():
        return src
    if src.is_file() and src.suffix.lower() == ".zip":
        dest = workdir / f"{label}_extract"
        dest.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(src) as zf:
            zf.extractall(dest)
        return dest
    raise SystemExit(f"Input '{src}' is neither a directory nor a .zip archive")


def find_markdown_files(root: Path):
    files = [p for p in sorted(root.rglob("*.md")) if p.is_file()]
    # Skip anything that looks like a macOS resource fork copy.
    return [p for p in files if not p.name.startswith("._")]


def find_reference_pdfs(root: Path):
    return [p for p in sorted(root.rglob("*.pdf"))
            if p.is_file() and not p.name.startswith("._")]


# ---------------------------------------------------------------------------
# Deterministic header / metadata extraction (section D)
# ---------------------------------------------------------------------------
META_LINE_RE = re.compile(r"^\*\*(?P<key>[^:*]+):\*\*\s*(?P<val>.*?)\s*$")
H1_RE = re.compile(r"^#\s+(?P<t>.+?)\s*$")
H2_RE = re.compile(r"^##\s+(?P<t>.+?)\s*$")
HR_RE = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")


def parse_yaml_front_matter(text: str):
    """Return (meta_dict, remaining_text) if YAML front matter is present."""
    if not text.startswith("---"):
        return None, text
    import yaml
    lines = text.splitlines(keepends=True)
    if not lines or lines[0].strip() != "---":
        return None, text
    for i in range(1, len(lines)):
        if lines[i].strip() in ("---", "..."):
            block = "".join(lines[1:i])
            rest = "".join(lines[i + 1:])
            try:
                data = yaml.safe_load(block) or {}
            except Exception:
                return None, text
            if isinstance(data, dict):
                return data, rest
            return None, text
    return None, text


def extract_header(text: str):
    """
    Deterministically split a source document into (metadata, body_markdown,
    warnings).

    Rule order (section D):
      1. YAML front matter is authoritative when present.
      2. Otherwise the first level-1 heading is the title and is removed from the
         body (only because it was used as the title).
      3. An H2 immediately following the title is the subtitle.
      4. Contiguous '**Key:** value' lines in the header block are metadata.
      5. A single trailing thematic break ('---') closing the header block is
         consumed. Nothing else is inferred; missing values are recorded, never
         invented.
    """
    warnings = []
    meta = {}
    ordered_meta = []  # (label, value) in authored order, for the title page

    fm, rest = parse_yaml_front_matter(text)
    if fm is not None:
        for k, v in fm.items():
            kl = str(k).strip().lower()
            meta[kl] = v
            ordered_meta.append((str(k).strip(), str(v)))
        body = rest
        # If YAML lacks a title, fall through to heading detection on the body.
        if "title" not in meta:
            meta2, body, w2 = _extract_prose_header(body)
            for k, v in meta2.items():
                meta.setdefault(k, v)
            warnings += w2
            ordered_meta = _merge_ordered(ordered_meta, meta2.get("_ordered", []))
        return meta, body, warnings

    return _extract_prose_header(text)


def _merge_ordered(a, b):
    seen = {k.lower() for k, _ in a}
    return a + [(k, v) for k, v in b if k.lower() not in seen]


def _extract_prose_header(text: str):
    warnings = []
    meta = {}
    ordered_meta = []
    lines = text.splitlines()
    i = 0
    n = len(lines)

    # (2) title = first H1
    while i < n and not lines[i].strip():
        i += 1
    title_consumed = False
    if i < n:
        m = H1_RE.match(lines[i])
        if m:
            meta["title"] = m.group("t").strip()
            i += 1
            title_consumed = True
        else:
            warnings.append("no level-1 heading found for title")

    # (3) subtitle = H2 immediately after the title
    j = i
    while j < n and not lines[j].strip():
        j += 1
    if title_consumed and j < n:
        m = H2_RE.match(lines[j])
        if m:
            meta["subtitle"] = m.group("t").strip()
            ordered_meta.append(("subtitle", meta["subtitle"]))
            i = j + 1

    # (4) contiguous '**Key:** value' metadata lines
    k = i
    while k < n and not lines[k].strip():
        k += 1
    consumed_any_meta = False
    while k < n:
        line = lines[k]
        if not line.strip():
            # allow a single blank line inside the metadata block only if more
            # metadata follows immediately
            look = k + 1
            while look < n and not lines[look].strip():
                look += 1
            if look < n and META_LINE_RE.match(lines[look].strip()):
                k = look
                continue
            break
        mm = META_LINE_RE.match(line.strip())
        if not mm:
            break
        key = mm.group("key").strip()
        val = mm.group("val").strip()
        meta[key.lower()] = val
        ordered_meta.append((key, val))
        consumed_any_meta = True
        k += 1
    if consumed_any_meta:
        i = k

    # (5) consume one closing thematic break
    p = i
    while p < n and not lines[p].strip():
        p += 1
    if p < n and HR_RE.match(lines[p]):
        i = p + 1

    body = "\n".join(lines[i:]).lstrip("\n")
    meta["_ordered"] = ordered_meta
    return meta, body, warnings


# ---------------------------------------------------------------------------
# Reference DOCX builder (section B / F)
# ---------------------------------------------------------------------------
def build_reference_docx(style: dict, family: str, out_path: Path):
    """
    Build reference/reference.docx carrying every named style the pipeline uses.
    Pandoc consumes this via --reference-doc; python-docx post-processing relies
    on the same named styles.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    st = style["styles"]
    page = style["page"]
    margins = style["margins"]

    doc = Document()

    def set_font(font, size=None, bold=None, italic=None, name=None, color=None):
        nm = name or family
        font.name = nm
        # Ensure east-asian / complex runs use the same face.
        rpr = font.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), nm)
        if size is not None:
            font.size = Pt(size)
        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic
        if color is not None:
            font.color.rgb = RGBColor(*color)

    # --- Normal / body ---
    normal = doc.styles["Normal"]
    b = st["body"]
    set_font(normal.font, size=b["size_pt"], bold=False, italic=False, color=(0, 0, 0))
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(b["line_spacing_pt"])
    pf.first_line_indent = Pt(b["first_line_indent_pt"])
    pf.space_before = Pt(b.get("space_before_pt", 0))
    pf.space_after = Pt(b.get("space_after_pt", 0))
    pf.widow_control = True

    # --- Headings ---
    def style_heading(word_style, spec, first_line_zero=True):
        s = doc.styles[word_style]
        set_font(s.font, size=spec["size_pt"], bold=spec.get("bold", True),
                 italic=spec.get("italic", False), color=(0, 0, 0))
        pf = s.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(spec.get("space_before_pt", 12))
        pf.space_after = Pt(spec.get("space_after_pt", 6))
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if first_line_zero:
            pf.first_line_indent = Pt(0)
        pf.keep_with_next = True
        pf.widow_control = True

    style_heading("Heading 1", st["heading_top"])
    style_heading("Heading 2", st["heading_section"])
    style_heading("Heading 3", st["heading_subsection"])
    style_heading("Heading 4", st["heading_subsubsection"])

    # --- Custom title-page styles ---
    def ensure_para_style(name):
        try:
            return doc.styles[name]
        except KeyError:
            return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # --- List / compact paragraph style (Pandoc styles tight-list items as
    #     "Compact"; it MUST exist without the body first-line indent, otherwise
    #     the +24pt first-line indent collapses the auto-number's slot and list
    #     numbers/bullets stop rendering). Body prose keeps its indent on Normal.
    compact = ensure_para_style("Compact")
    compact.base_style = doc.styles["Normal"]
    set_font(compact.font, size=b["size_pt"], color=(0, 0, 0))
    cpf = compact.paragraph_format
    cpf.first_line_indent = Pt(0)
    cpf.left_indent = Pt(0)
    cpf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cpf.line_spacing = Pt(b["line_spacing_pt"])
    cpf.space_before = Pt(0)
    cpf.space_after = Pt(0)

    def make_titlepage_style(name, spec):
        s = ensure_para_style(name)
        s.base_style = doc.styles["Normal"]
        set_font(s.font, size=spec["size_pt"], bold=spec.get("bold", False),
                 italic=spec.get("italic", False), color=(0, 0, 0))
        pf = s.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return s

    make_titlepage_style("MWE Title", st["title"])
    make_titlepage_style("MWE Subtitle", st["subtitle"])
    make_titlepage_style("MWE Meta", st["meta_line"])
    make_titlepage_style("MWE Author", st["author"])
    make_titlepage_style("MWE Affiliation", st["affiliation"])
    spacer = make_titlepage_style("MWE Spacer", {"size_pt": st["body"]["size_pt"]})

    # --- Block quote (pandoc: "Block Text") ---
    bq = ensure_para_style("Block Text")
    bq.base_style = doc.styles["Normal"]
    q = st.get("blockquote", {})
    set_font(bq.font, size=q.get("size_pt", 12), italic=q.get("italic", False), color=(0, 0, 0))
    bq.paragraph_format.left_indent = Pt(q.get("left_indent_pt", 24))
    bq.paragraph_format.first_line_indent = Pt(0)

    # --- Source code (pandoc: "Source Code") ---
    sc = ensure_para_style("Source Code")
    sc.base_style = doc.styles["Normal"]
    c = st.get("code", {})
    set_font(sc.font, size=c.get("size_pt", 10.5), name=c.get("family", "DejaVu Sans Mono"), color=(0, 0, 0))
    sc.paragraph_format.first_line_indent = Pt(0)
    sc.paragraph_format.left_indent = Pt(c.get("left_indent_pt", 12))
    sc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # inline verbatim char style
    try:
        vch = doc.styles["Verbatim Char"]
    except KeyError:
        vch = doc.styles.add_style("Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    set_font(vch.font, size=c.get("size_pt", 10.5), name=c.get("family", "DejaVu Sans Mono"), color=(0, 0, 0))

    # --- Caption ---
    cap = ensure_para_style("Image Caption")
    cap.base_style = doc.styles["Normal"]
    cp = st.get("caption", {})
    set_font(cap.font, size=cp.get("size_pt", 10.5), italic=cp.get("italic", True), color=(0, 0, 0))
    cap.paragraph_format.first_line_indent = Pt(0)

    # --- References (hanging indent) ---
    ref = ensure_para_style("MWE References")
    ref.base_style = doc.styles["Normal"]
    rf = st.get("references", {})
    set_font(ref.font, size=rf.get("size_pt", 12), color=(0, 0, 0))
    rpf = ref.paragraph_format
    rpf.left_indent = Pt(rf.get("hanging_indent_pt", 24))
    rpf.first_line_indent = Pt(-rf.get("hanging_indent_pt", 24))
    rpf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    rpf.line_spacing = Pt(rf.get("line_spacing_pt", 27))

    # --- Page geometry on the single default section ---
    sec = doc.sections[0]
    sec.page_width = Pt(page["width_pt"])
    sec.page_height = Pt(page["height_pt"])
    sec.top_margin = Pt(margins["top_pt"])
    sec.bottom_margin = Pt(margins["bottom_pt"])
    sec.left_margin = Pt(margins["left_pt"])
    sec.right_margin = Pt(margins["right_pt"])
    sec.header_distance = Pt(35)
    sec.footer_distance = Pt(35)

    # Reference doc must have no stray body content; a single empty paragraph is
    # fine and pandoc ignores it.
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Markdown -> DOCX via Pandoc (section E)
# ---------------------------------------------------------------------------
def markdown_to_docx(body_md: str, reference_docx: Path, out_docx: Path, workdir: Path):
    """
    Convert body Markdown to DOCX with Pandoc. Smart typography is DISABLED so
    no character (quote, dash, symbol) is altered. The body is passed verbatim.
    """
    tmp_md = workdir / (out_docx.stem + ".body.md")
    tmp_md.write_text(body_md, encoding="utf-8")
    # markdown-smart: read as pandoc markdown but do NOT apply smart typography.
    reader = "markdown-smart+fenced_code_attributes+pipe_tables+footnotes+tex_math_dollars"
    cmd = [
        "pandoc",
        str(tmp_md),
        "--from", reader,
        "--to", "docx",
        "--reference-doc", str(reference_docx),
        "--wrap", "preserve",
        "-o", str(out_docx),
    ]
    rc, out, err = run(cmd)
    if rc != 0:
        raise RuntimeError(f"pandoc failed: {err.strip()}")
    return out_docx


def markdown_to_plain(md_text: str, workdir: Path, tag: str) -> str:
    """Normalised plaintext rendering of Markdown for the semantic check.

    Uses Pandoc's own parser (not regexes) with smart typography disabled and no
    wrapping, so characters are preserved exactly.
    """
    tmp_md = workdir / f"{tag}.norm.md"
    tmp_md.write_text(md_text, encoding="utf-8")
    cmd = ["pandoc", str(tmp_md), "--from",
           "markdown-smart+pipe_tables+footnotes", "--to", "plain",
           "--wrap", "none"]
    rc, out, err = run(cmd)
    if rc != 0:
        raise RuntimeError(f"pandoc plain failed: {err.strip()}")
    return out


# ---------------------------------------------------------------------------
# python-docx post-processing (sections F)
# ---------------------------------------------------------------------------
def _add_titlepage_paragraph(anchor, text, style_name, runs_spec=None):
    """Insert a styled paragraph *before* `anchor`. Returns the new paragraph."""
    p = anchor.insert_paragraph_before(style=style_name)
    if runs_spec is None:
        if text:
            p.add_run(text)
    else:
        for chunk, bold, italic in runs_spec:
            r = p.add_run(chunk)
            if bold is not None:
                r.bold = bold
            if italic is not None:
                r.italic = italic
    return p


def _add_spacer(anchor, pt):
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    p = anchor.insert_paragraph_before(style="MWE Spacer")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    return p


def build_title_page_and_postprocess(docx_path: Path, meta: dict, style: dict,
                                     family: str):
    """
    Insert a dedicated title page (section F) followed by a page break, set the
    page geometry, sanitise document metadata, and apply a hanging indent to a
    References section if one exists. Purely presentational; no body text is
    altered.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(str(docx_path))

    # ---- page geometry (belt-and-braces; reference doc already sets it) ----
    page = style["page"]
    margins = style["margins"]
    for sec in doc.sections:
        sec.page_width = Pt(page["width_pt"])
        sec.page_height = Pt(page["height_pt"])
        sec.top_margin = Pt(margins["top_pt"])
        sec.bottom_margin = Pt(margins["bottom_pt"])
        sec.left_margin = Pt(margins["left_pt"])
        sec.right_margin = Pt(margins["right_pt"])

    tp = style["title_page"]

    # Determine an anchor: the first body paragraph. If the document is empty,
    # add one so we have something to anchor to.
    if not doc.paragraphs:
        doc.add_paragraph("")
    anchor = doc.paragraphs[0]

    # We insert title-page content *before* the anchor, top-down, then a page
    # break paragraph, so the body starts on page 2.

    # 1) top whitespace
    _add_spacer(anchor, tp.get("title_top_offset_pt", 275))

    # 2) title
    title = str(meta.get("title", "")).strip()
    if title:
        _add_titlepage_paragraph(anchor, title, "MWE Title")

    # 3) subtitle
    subtitle = str(meta.get("subtitle", "")).strip()
    if subtitle:
        _add_spacer(anchor, tp.get("gap_title_to_subtitle_pt", 6))
        _add_titlepage_paragraph(anchor, subtitle, "MWE Subtitle")

    # 4) metadata block: every authored '**Key:** value' line EXCEPT author /
    #    affiliation (rendered separately near the bottom), in authored order.
    ordered = meta.get("_ordered", [])
    meta_block = [(k, v) for (k, v) in ordered
                  if k.lower() not in ("author", "affiliation", "subtitle")]
    if meta_block:
        _add_spacer(anchor, tp.get("gap_subtitle_to_meta_pt", 18))
        for k, v in meta_block:
            _add_titlepage_paragraph(
                anchor, None, "MWE Meta",
                runs_spec=[(f"{k}: ", True, True), (v, False, True)])

    # 5) author block near the bottom
    author = str(meta.get("author", "")).strip()
    affiliation = str(meta.get("affiliation", "")).strip()
    if author or affiliation:
        _add_spacer(anchor, tp.get("author_block_from_bottom_pt", 140))
    if author:
        _add_titlepage_paragraph(anchor, author, "MWE Author")
    if affiliation:
        _add_spacer(anchor, tp.get("gap_author_to_affiliation_pt", 12))
        _add_titlepage_paragraph(anchor, affiliation, "MWE Affiliation")

    # 6) page break so the body begins on its own page
    brk = anchor.insert_paragraph_before(style="MWE Spacer")
    run_ = brk.add_run()
    run_.add_break(WD_BREAK.PAGE)

    # ---- References hanging indent (section A), only if a section exists ----
    _apply_references_style(doc)

    # ---- metadata sanitisation (section F) ----
    cp = doc.core_properties
    cp.title = title or ""
    cp.author = str(meta.get("author", "") or "")
    cp.last_modified_by = ""
    cp.comments = ""
    cp.category = ""
    cp.keywords = ""
    cp.subject = ""
    cp.identifier = ""
    cp.revision = 1

    doc.save(str(docx_path))
    _strip_empty_comment_parts(docx_path)


def _strip_empty_comment_parts(docx_path: Path):
    """Remove empty comments/tracked-change parts so the deliverable literally
    contains no comment parts (section F: 'no comments or tracked changes')."""
    import zipfile
    tmp = docx_path.with_suffix(".clean.docx")
    drop_prefixes = ("word/comments",)  # comments.xml, commentsExtended.xml, ...
    with zipfile.ZipFile(docx_path) as zin:
        keep = {}
        remove = set()
        for item in zin.namelist():
            data = zin.read(item)
            base = item.rsplit("/", 1)[-1]
            if item.startswith(drop_prefixes) and b"<w:comment " not in data:
                remove.add(base)
                continue
            keep[item] = data
    if not remove:
        return  # nothing to strip
    # Also scrub relationships and content-type overrides that referenced them.
    import re
    for name, data in list(keep.items()):
        if name.endswith(".rels"):
            txt = data.decode("utf-8")
            txt = re.sub(r"<Relationship\b[^>]*Target=\"[^\"]*comments[^\"]*\"[^>]*/>",
                         "", txt)
            keep[name] = txt.encode("utf-8")
        if name == "[Content_Types].xml":
            txt = data.decode("utf-8")
            txt = re.sub(r"<Override\b[^>]*PartName=\"/word/comments[^\"]*\"[^>]*/>",
                         "", txt)
            keep[name] = txt.encode("utf-8")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in keep.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


def _apply_references_style(doc):
    """If a heading named References/Bibliography/Works Cited exists, apply the
    hanging-indent style to the paragraphs beneath it up to the next heading."""
    heads = {"references", "bibliography", "works cited"}
    in_refs = False
    for p in doc.paragraphs:
        sname = (p.style.name or "").lower()
        if sname.startswith("heading"):
            in_refs = p.text.strip().lower() in heads
            continue
        if in_refs and p.text.strip():
            p.style = doc.styles["MWE References"]


# ---------------------------------------------------------------------------
# DOCX -> PDF (section G)
# ---------------------------------------------------------------------------
def docx_to_pdf(docx_path: Path, out_dir: Path, profile_dir: Path) -> Path:
    soffice = which("libreoffice") or which("soffice")
    env = dict(os.environ)
    # Isolate LibreOffice under a private, writable profile so runs are
    # deterministic and carry no environment personalisation.
    env["HOME"] = str(profile_dir)
    (profile_dir / "cache").mkdir(parents=True, exist_ok=True)
    (profile_dir / "config").mkdir(parents=True, exist_ok=True)
    env["XDG_CACHE_HOME"] = str(profile_dir / "cache")
    env["XDG_CONFIG_HOME"] = str(profile_dir / "config")
    env["SAL_USE_VCLPLUGIN"] = "svp"
    cmd = [
        soffice, "--headless", "--norestore", "--nolockcheck",
        f"-env:UserInstallation=file://{profile_dir}",
        "--convert-to", "pdf:writer_pdf_Export",
        "--outdir", str(out_dir), str(docx_path),
    ]
    rc, out, err = run(cmd, env=env)
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if rc != 0 or not pdf_path.exists():
        raise RuntimeError(f"LibreOffice PDF conversion failed: {err.strip() or out.strip()}")
    return pdf_path


# ---------------------------------------------------------------------------
# Text extraction + normalisation for the semantic check (section I)
# ---------------------------------------------------------------------------
# A "thematic break" line: only -, _ or * (>=3). Pandoc's plain writer renders
# a Markdown horizontal rule as a full line of dashes; DOCX/PDF render it as an
# empty paragraph. These carry no textual content and are excluded.
HR_LINE_RE = re.compile(r"^\s*([-_*]\s*){3,}$")
# A standalone enumerator token, e.g. "1." "12)" — how ordered-list numbers and
# heading numbers appear once list markers are separated by the renderer.
ENUM_TOKEN_RE = re.compile(r"^\d+[.)]$")
# A leading list/enumerator marker at the start of a line.
LEAD_MARKER_RE = re.compile(r"^(\s*)([-*+•]|\d+[.)])\s+")
# A line consisting solely of an enumerator (e.g. "1." on its own line), which
# is how some PDF text extractors place ordered-list numbers. Note this does not
# match "5.1" (no trailing separator+space), so subsection numbers are kept.
ENUM_ONLY_LINE_RE = re.compile(r"^\s*\d+[.)]\s*$")
# Pure punctuation / rule fragments to ignore as non-content.
PUNCT_TOKEN_RE = re.compile(r"^[\-–—_*=.,;:]+$")


def normalize_tokens(text: str):
    """
    Turn rendered text into a comparable content-token sequence.

    Excludes only presentation artefacts (section I item 4): thematic-break rule
    lines, list/enumerator markers, purely presentational whitespace, and
    line-wrap hyphenation. Preserves every word, punctuation mark inside words,
    Unicode symbol, capitalisation, and section-heading TEXT.
    """
    # Collect tokens, tracking which token ends a line (a wrap point).
    raw_tokens = []  # (token, is_line_end)
    for line in text.splitlines():
        if not line.strip():
            continue
        if HR_LINE_RE.match(line):
            continue  # thematic break -> no content
        line = LEAD_MARKER_RE.sub(r"\1", line)  # drop a leading bullet/number
        parts = line.split()
        for j, p in enumerate(parts):
            raw_tokens.append((p, j == len(parts) - 1))

    # De-hyphenate line-wrap splits ONLY across a line break: a token that ends a
    # line with a trailing '-' is an authored hyphenated word broken by wrapping;
    # rejoin it with the first token of the next line, KEEPING the hyphen (Word/
    # LibreOffice auto-hyphenation is off, so no hyphen is renderer-inserted).
    merged = []
    i = 0
    while i < len(raw_tokens):
        tok, is_end = raw_tokens[i]
        if (is_end and tok.endswith("-") and not ENUM_TOKEN_RE.match(tok)
                and i + 1 < len(raw_tokens)):
            merged.append(tok + raw_tokens[i + 1][0])
            i += 2
        else:
            merged.append(tok)
            i += 1

    out = []
    for tok in merged:
        if ENUM_TOKEN_RE.match(tok):
            continue  # standalone list/heading enumerator
        if PUNCT_TOKEN_RE.match(tok):
            continue  # rule fragment / stray punctuation run
        out.append(tok)
    return out


def docx_to_plain(docx_path: Path) -> str:
    """Render a DOCX to plain text with the SAME pandoc writer used for the
    source, so horizontal rules, list numbers and headings normalise
    identically on both sides of the comparison."""
    cmd = ["pandoc", str(docx_path), "--from", "docx", "--to", "plain",
           "--wrap", "none"]
    rc, out, err = run(cmd)
    if rc != 0:
        raise RuntimeError(f"pandoc docx->plain failed: {err.strip()}")
    return out


def extract_pdf_text(pdf_path: Path):
    """Return (full_text, page_count, body_text). Body = pages 2+ (page 1 is the
    title page)."""
    import fitz
    d = fitz.open(str(pdf_path))
    parts = [pg.get_text("text") for pg in d]
    n = len(d)
    d.close()
    full = "\n".join(parts)
    body = "\n".join(parts[1:]) if n > 1 else full
    return full, n, body


def multiset_diff(expected, actual):
    ce, ca = Counter(expected), Counter(actual)
    return ce - ca, ca - ce


def semantic_check_body(source_body_md, body_docx_path, workdir, tag):
    """
    Source-fidelity gate. Compares the source Markdown body against the
    body-only DOCX (BEFORE the title page is inserted) using pandoc's plain
    writer on both sides. Ordered comparison catches missing sentences, altered
    punctuation, changed Unicode symbols, reordered paragraphs, altered
    headings, and missing list items.
    """
    import difflib
    result = {"passed": True, "details": [], "warnings": []}

    src_tokens = normalize_tokens(markdown_to_plain(source_body_md, workdir, tag))
    docx_tokens = normalize_tokens(docx_to_plain(body_docx_path))

    miss, add = multiset_diff(src_tokens, docx_tokens)
    if miss or add:
        result["passed"] = False
        result["details"].append(
            f"source<->DOCX content mismatch: missing={_sample(miss)} added={_sample(add)}")
    else:
        sm = difflib.SequenceMatcher(a=src_tokens, b=docx_tokens)
        if sm.ratio() < 0.999:
            result["passed"] = False
            result["details"].append(
                f"source<->DOCX reordering detected (ratio={sm.ratio():.4f})")
    return result


def content_charstream(text: str) -> str:
    """
    Whitespace-insensitive content signature: drop thematic-break rule lines and
    the leading list/enumerator marker of each line, then concatenate every
    remaining non-whitespace character in order.

    This is the correct invariant for 'the PDF faithfully renders the DOCX': line
    wrapping, hyphenation at existing dashes, and PDF-extractor spacing quirks
    change only whitespace, never actual characters — so they vanish here, while
    any genuinely missing/added/changed character (a dropped sentence, an altered
    Unicode symbol, changed punctuation) still shows up. Leading enumerators
    (ordered-list and section-heading numbers) are dropped on BOTH sides because
    Pandoc's plain writer emits them as text whereas the PDF's numbering is not
    part of the extractable text layer; they are preserved verbatim by
    construction and confirmed in the visual QA renders.
    """
    kept = []
    for line in text.splitlines():
        if not line.strip():
            continue
        if HR_LINE_RE.match(line):
            continue
        if ENUM_ONLY_LINE_RE.match(line):
            continue  # a list number extracted onto its own line
        line = LEAD_MARKER_RE.sub(r"\1", line)
        kept.append(line)
    joined = "".join(kept)
    return "".join(ch for ch in joined if not ch.isspace())


def semantic_check_pdf(body_docx_path, pdf_body_text):
    """
    PDF-fidelity check. Confirms the PDF body (pages 2+) carries exactly the same
    characters as the body-only DOCX (selectable text, no dropped words, no
    altered symbols), tolerant of line-wrap and extractor spacing only. The title
    page (page 1) is excluded on both sides.
    """
    result = {"passed": True, "details": [], "warnings": []}
    docx_stream = content_charstream(docx_to_plain(body_docx_path))
    pdf_stream = content_charstream(pdf_body_text)
    if docx_stream != pdf_stream:
        result["passed"] = False
        result["details"].append(_charstream_diff("DOCX", docx_stream,
                                                   "PDF", pdf_stream))
    return result


def _charstream_diff(na, a, nb, b):
    import difflib
    sm = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag != "equal":
            ctx = 25
            return (f"{na}<->{nb} character mismatch near "
                    f"{na}[...{a[max(0,i1-ctx):i2+ctx]!r}...] vs "
                    f"{nb}[...{b[max(0,j1-ctx):j2+ctx]!r}...] "
                    f"(ratio={sm.ratio():.5f})")
    return f"{na}<->{nb} mismatch (ratio={sm.ratio():.5f})"


def titlepage_present_in_pdf(meta, pdf_full_text):
    """Presence check (warning-level): every title-page metadata word should
    appear somewhere in the PDF."""
    tp_tokens = [t for t in normalize_tokens(title_page_plaintext(meta)) if _is_wordish(t)]
    pdf_tokens = Counter(t for t in normalize_tokens(pdf_full_text) if _is_wordish(t))
    missing = Counter(tp_tokens) - pdf_tokens
    return list(missing.elements())


def _is_wordish(tok):
    return any(ch.isalnum() for ch in tok)


def _sample(counter, k=12):
    items = list(counter.elements())
    return items[:k] + (["...(+%d)" % (len(items) - k)] if len(items) > k else [])


# ---------------------------------------------------------------------------
# Visual verification (section J)
# ---------------------------------------------------------------------------
def render_pdf_pages(pdf_path: Path, out_dir: Path, zoom=1.6):
    import fitz
    out_dir.mkdir(parents=True, exist_ok=True)
    d = fitz.open(str(pdf_path))
    imgs = []
    for i, pg in enumerate(d):
        pix = pg.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        p = out_dir / f"page_{i+1:03d}.png"
        pix.save(str(p))
        imgs.append(p)
    d.close()
    return imgs


def audit_pdf(pdf_path: Path):
    """Automated page audit: blank pages, missing glyphs, clipping heuristics."""
    import fitz
    findings = []
    d = fitz.open(str(pdf_path))
    W = d[0].rect.width if len(d) else 0
    H = d[0].rect.height if len(d) else 0
    for i, pg in enumerate(d):
        text = pg.get_text("text").strip()
        # blank page (allow a genuinely short last page)
        if not text and pg.get_images() == [] and pg.get_drawings() == []:
            findings.append(f"page {i+1}: blank")
        # missing glyph markers
        rd = pg.get_text("rawdict")
        for blk in rd.get("blocks", []):
            for ln in blk.get("lines", []):
                for sp in ln.get("spans", []):
                    for ch in sp.get("chars", []):
                        c = ch.get("c")
                        if c == "�" or c == "\x00":
                            findings.append(f"page {i+1}: replacement/missing glyph")
                            break
        # clipping: text bbox outside the printable area
        for blk in pg.get_text("blocks"):
            x0, y0, x1, y1 = blk[:4]
            if x0 < -1 or y0 < -1 or x1 > W + 1 or y1 > H + 1:
                findings.append(f"page {i+1}: content outside page bounds")
                break
    d.close()
    # de-duplicate while preserving order
    seen = set()
    uniq = []
    for f in findings:
        if f not in seen:
            uniq.append(f); seen.add(f)
    return uniq


def make_contact_sheet(entries, out_path: Path, cols=6, thumb_w=300):
    """entries: list of (label, image_path). Builds one PNG grid."""
    from PIL import Image, ImageDraw
    if not entries:
        return None
    thumbs = []
    for label, img in entries:
        try:
            im = Image.open(img).convert("RGB")
        except Exception:
            continue
        ratio = thumb_w / im.width
        im = im.resize((thumb_w, int(im.height * ratio)))
        thumbs.append((label, im))
    if not thumbs:
        return None
    th = max(im.height for _, im in thumbs)
    pad, labelh = 8, 16
    cell_w = thumb_w + pad
    cell_h = th + pad + labelh
    rows = (len(thumbs) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * cell_w + pad, rows * cell_h + pad), "white")
    draw = ImageDraw.Draw(sheet)
    for idx, (label, im) in enumerate(thumbs):
        r, c = divmod(idx, cols)
        x = pad + c * cell_w
        y = pad + r * cell_h
        draw.text((x, y), label[:46], fill="black")
        sheet.paste(im, (x, y + labelh))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Main orchestration
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="Batch Markdown -> DOCX/PDF for OSF.")
    ap.add_argument("--input", required=True, help="Markdown .zip or directory")
    ap.add_argument("--reference", required=True, help="Reference PDF .zip or directory")
    ap.add_argument("--output", required=True, help="Output directory (osf-export)")
    ap.add_argument("--config", default=str(DEFAULT_CONFIG))
    ap.add_argument("--reference-docx", default=str(DEFAULT_REFERENCE_DOCX))
    ap.add_argument("--keep-build", action="store_true",
                    help="keep the temporary build directory")
    args = ap.parse_args()

    started = _dt.datetime.now()
    log(f"[batch_export] started {started.isoformat(timespec='seconds')}")

    # --- dependencies ---
    problems = check_dependencies()
    if problems:
        log(f"::: {DEPENDENCY_REQUIREMENT_BLOCKED} :::")
        for p in problems:
            log(f"  - {p}")
        raise SystemExit(3)

    style = load_style(Path(args.config))

    # --- font boundary ---
    family, font_status = resolve_font(style)
    log(f"[batch_export] effective font family: {family} "
        f"(primary '{font_status['primary_family']}' "
        f"installed={font_status['primary_installed']})")

    out_root = Path(args.output).resolve()
    dirs = {k: out_root / k for k in
            ("docx", "pdf", "qa-renders", "reports", "reference", "config")}
    for d in dirs.values():
        d.mkdir(parents=True, exist_ok=True)
    (dirs["reports"] / "failed").mkdir(exist_ok=True)

    build_dir = Path(tempfile.mkdtemp(prefix="mwe_build_"))
    lo_profile = build_dir / "lo_profile"
    lo_profile.mkdir(parents=True, exist_ok=True)

    try:
        md_root = materialise_input(Path(args.input).resolve(), build_dir, "input")
        ref_root = materialise_input(Path(args.reference).resolve(), build_dir, "reference")

        md_files = find_markdown_files(md_root)
        ref_pdfs = find_reference_pdfs(ref_root)
        log(f"[batch_export] {len(md_files)} markdown source(s), "
            f"{len(ref_pdfs)} reference PDF(s)")

        # --- reference.docx ---
        reference_docx = Path(args.reference_docx).resolve()
        build_reference_docx(style, family, reference_docx)
        shutil.copy2(reference_docx, dirs["reference"] / "reference.docx")
        shutil.copy2(Path(args.config), dirs["config"] / Path(args.config).name)
        log(f"[batch_export] built reference DOCX -> {reference_docx}")

        rows = []
        summary = {
            "found": len(md_files), "converted": 0, "failed": 0, "warnings": 0,
            "missing_metadata": [], "unsupported_features": [],
            "layout_warnings": [], "human_review": [], "semantic_failed": [],
        }
        contact_entries = {"page1": [], "firstbody": [], "last": []}

        for md in md_files:
            row = _process_one(md, style, family, reference_docx, dirs,
                               build_dir, lo_profile, summary, contact_entries)
            rows.append(row)

        # --- manifest ---
        _write_manifest(dirs["reports"] / "conversion-manifest.csv", rows)

        # --- contact sheets ---
        cs_dir = dirs["qa-renders"] / "_contact_sheets"
        make_contact_sheet(contact_entries["page1"], cs_dir / "contact_titlepages.png")
        make_contact_sheet(contact_entries["firstbody"], cs_dir / "contact_firstbody.png")
        make_contact_sheet(contact_entries["last"], cs_dir / "contact_lastpage.png")

        # --- classification ---
        classification = _classify(summary)
        _write_summary(dirs["reports"] / "conversion-summary.md", summary,
                       font_status, len(ref_pdfs), classification, family, started)

        log("")
        log(f"[batch_export] classification: {classification}")
        log(f"[batch_export] converted={summary['converted']} "
            f"failed={summary['failed']} warnings={summary['warnings']}")
        log(f"[batch_export] outputs under {out_root}")

    finally:
        if not args.keep_build:
            shutil.rmtree(build_dir, ignore_errors=True)
        else:
            log(f"[batch_export] build dir kept: {build_dir}")


def _process_one(md, style, family, reference_docx, dirs, build_dir,
                 lo_profile, summary, contact_entries):
    name = md.stem
    row = {c: "" for c in MANIFEST_COLUMNS}
    row["source_markdown"] = md.name
    warnings = []
    log(f"\n[convert] {md.name}")

    try:
        raw = md.read_text(encoding="utf-8")
        meta, body_md, hwarn = extract_header(raw)
        warnings += hwarn

        # metadata for manifest
        for key in ("title", "subtitle", "version", "date", "status",
                    "author", "affiliation"):
            row[key] = str(meta.get(key, "") or "")

        # record missing metadata (section D rule 6)
        missing_meta = [k for k in ("title", "version", "date", "status", "author")
                        if not meta.get(k)]
        if missing_meta:
            summary["missing_metadata"].append(f"{md.name}: {', '.join(missing_meta)}")
            warnings.append("missing metadata: " + ", ".join(missing_meta))

        stem = str(meta.get("output_stem") or name)
        docx_out = dirs["docx"] / f"{stem}.docx"
        pdf_out = dirs["pdf"] / f"{stem}.pdf"

        # 1) markdown -> body-only DOCX (retained for semantic checks)
        body_docx = build_dir / f"{stem}.body.docx"
        markdown_to_docx(body_md, reference_docx, body_docx, build_dir)

        # 1a) source-fidelity gate on the body-only DOCX (before title page)
        sem_body = semantic_check_body(body_md, body_docx, build_dir, stem)

        # 2) title page + post-processing on the deliverable DOCX
        shutil.copy2(body_docx, docx_out)
        build_title_page_and_postprocess(docx_out, meta, style, family)
        docx_render_pass = docx_out.exists() and docx_out.stat().st_size > 0

        # 3) docx -> pdf
        pdf_tmp_dir = build_dir / "pdf_out"
        pdf_tmp_dir.mkdir(exist_ok=True)
        pdf_made = docx_to_pdf(docx_out, pdf_tmp_dir, lo_profile)
        shutil.move(str(pdf_made), str(pdf_out))

        # 4) PDF-fidelity check (body vs body) + title-page presence
        pdf_text, pdf_pages, pdf_body_text = extract_pdf_text(pdf_out)
        sem_pdf = semantic_check_pdf(body_docx, pdf_body_text)
        tp_missing = titlepage_present_in_pdf(meta, pdf_text)
        if tp_missing:
            warnings.append(f"title-page tokens not found in PDF: {tp_missing[:8]}")

        sem = {"passed": sem_body["passed"] and sem_pdf["passed"],
               "details": sem_body["details"] + sem_pdf["details"],
               "warnings": sem_body["warnings"] + sem_pdf["warnings"]}
        warnings += sem["warnings"]

        # 5) visual audit + renders
        qa_dir = dirs["qa-renders"] / stem
        imgs = render_pdf_pages(pdf_out, qa_dir)
        audit = audit_pdf(pdf_out)
        if audit:
            summary["layout_warnings"].append(f"{md.name}: {'; '.join(audit)}")
            warnings += audit
        pdf_render_pass = bool(imgs) and not any(
            a for a in audit if ("blank" in a or "missing glyph" in a
                                 or "outside page" in a))

        # contact-sheet entries
        if imgs:
            contact_entries["page1"].append((stem, imgs[0]))
            contact_entries["firstbody"].append((stem, imgs[1] if len(imgs) > 1 else imgs[0]))
            contact_entries["last"].append((stem, imgs[-1]))

        # sizes + hashes
        row["output_docx"] = docx_out.name
        row["output_pdf"] = pdf_out.name
        row["source_bytes"] = md.stat().st_size
        row["docx_bytes"] = docx_out.stat().st_size
        row["pdf_bytes"] = pdf_out.stat().st_size
        row["pdf_pages"] = pdf_pages
        row["docx_render_pass"] = "PASS" if docx_render_pass else "FAIL"
        row["pdf_render_pass"] = "PASS" if pdf_render_pass else "FAIL"
        row["semantic_check"] = "PASS" if sem["passed"] else "FAIL"
        row["warnings"] = " | ".join(warnings)
        row["sha256_source"] = sha256_of(md)
        row["sha256_docx"] = sha256_of(docx_out)
        row["sha256_pdf"] = sha256_of(pdf_out)

        # classification bookkeeping
        if not sem["passed"]:
            summary["semantic_failed"].append(md.name)
            # isolate failed outputs under reports/failed
            fdir = dirs["reports"] / "failed"
            shutil.copy2(docx_out, fdir / docx_out.name)
            shutil.copy2(pdf_out, fdir / pdf_out.name)
            (fdir / f"{stem}.semantic.txt").write_text(
                "\n".join(sem["details"]), encoding="utf-8")
            summary["human_review"].append(f"{md.name}: SEMANTIC_EQUIVALENCE_FAILED")
        if not pdf_render_pass:
            summary["human_review"].append(f"{md.name}: LAYOUT_VERIFICATION_FAILED")

        if sem["passed"] and docx_render_pass and pdf_render_pass:
            summary["converted"] += 1
        else:
            summary["failed"] += 1
        if warnings:
            summary["warnings"] += 1
        log(f"[convert] {md.name}: docx={row['docx_render_pass']} "
            f"pdf={row['pdf_render_pass']} semantic={row['semantic_check']} "
            f"pages={pdf_pages}")

    except Exception as exc:  # isolate a failed file, continue the batch
        summary["failed"] += 1
        row["docx_render_pass"] = row["docx_render_pass"] or "FAIL"
        row["pdf_render_pass"] = row["pdf_render_pass"] or "FAIL"
        row["semantic_check"] = row["semantic_check"] or "ERROR"
        row["warnings"] = f"ERROR: {exc}"
        summary["human_review"].append(f"{md.name}: ERROR {exc}")
        log(f"[convert] {md.name}: ERROR {exc}")
    return row


def title_page_plaintext(meta):
    """Plain-text rendering of the title page, mirroring exactly what
    build_title_page_and_postprocess writes (subtitle/author/affiliation without
    labels; other metadata as 'Key: value')."""
    parts = []
    if meta.get("title"):
        parts.append(str(meta["title"]))
    if meta.get("subtitle"):
        parts.append(str(meta["subtitle"]))
    for k, v in meta.get("_ordered", []):
        if k.lower() in ("author", "affiliation", "subtitle"):
            continue
        parts.append(f"{k}: {v}")
    if meta.get("author"):
        parts.append(str(meta["author"]))
    if meta.get("affiliation"):
        parts.append(str(meta["affiliation"]))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Manifest + summary writers
# ---------------------------------------------------------------------------
MANIFEST_COLUMNS = [
    "source_markdown", "output_docx", "output_pdf",
    "title", "subtitle", "version", "date", "status", "author", "affiliation",
    "source_bytes", "docx_bytes", "pdf_bytes", "pdf_pages",
    "docx_render_pass", "pdf_render_pass", "semantic_check", "warnings",
    "sha256_source", "sha256_docx", "sha256_pdf",
]


def _write_manifest(path: Path, rows):
    with open(path, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=MANIFEST_COLUMNS)
        w.writeheader()
        for r in rows:
            w.writerow({c: r.get(c, "") for c in MANIFEST_COLUMNS})


def _classify(summary):
    if summary["semantic_failed"]:
        return SEMANTIC_EQUIVALENCE_FAILED
    if summary["failed"] and summary["converted"] == 0:
        return LAYOUT_VERIFICATION_FAILED
    if summary["failed"] or summary["warnings"] or summary["missing_metadata"]:
        return BATCH_EXPORT_COMPLETE_WITH_WARNINGS
    return BATCH_EXPORT_COMPLETE


def _write_summary(path, summary, font_status, n_ref_pdfs, classification,
                   family, started):
    def block(title, items):
        if not items:
            return f"- **{title}:** none\n"
        return f"- **{title}:**\n" + "".join(f"    - {i}\n" for i in items)

    txt = f"""# Conversion Summary

- **Primary classification:** {classification}
- **Generated:** {_dt.datetime.now().isoformat(timespec='seconds')}
- **Started:** {started.isoformat(timespec='seconds')}

## Font status (section C)

- Intended primary font: `{font_status['primary_family']}` — installed: **{font_status['primary_installed']}**
- Substitute family: `{font_status['substitute_family']}` — approved: **{font_status['substitute_approved']}**, installed: **{font_status['substitute_installed']}**
- Effective font applied to all styles: **{family}**

## Counts

- Markdown files found: **{summary['found']}**
- Successfully converted (DOCX + PDF + semantic all PASS): **{summary['converted']}**
- Blocked / failed: **{summary['failed']}**
- Files with warnings: **{summary['warnings']}**
- Reference PDFs analysed: **{n_ref_pdfs}**

## Diagnostics

{block("Missing metadata", summary['missing_metadata'])}
{block("Unsupported Markdown features", summary['unsupported_features'])}
{block("Files with layout warnings", summary['layout_warnings'])}
{block("Files requiring human review", summary['human_review'])}
{block("Semantic-equivalence failures", summary['semantic_failed'])}

## No-semantic-change audit

Every document's Markdown body was passed to Pandoc verbatim with smart
typography disabled; no source text was paraphrased, reordered, translated, or
renumbered. The normalised token content of each source was compared against its
DOCX and PDF. Result: {"all documents PASS" if not summary['semantic_failed'] else "FAILURES PRESENT — see reports/failed/"}.

_No OSF upload was performed. Source Markdown files were not modified._
"""
    path.write_text(txt, encoding="utf-8")


if __name__ == "__main__":
    main()
